from django.shortcuts import render, redirect
from django.core.files.storage import FileSystemStorage
from django.contrib.auth.decorators import login_required
from django.http import JsonResponse
from .models import UserActivity
from django.conf import settings
import json
import os
import pdfplumber
import docx
import PyPDF2
import re

# --- MAIN NAVIGATION VIEWS ---

@login_required
def dashboard_view(request):
    # Fetch the 5 most recent activities for this user
    # The minus sign in '-timestamp' puts the most recent files at the top
    recent_activities = UserActivity.objects.filter(user=request.user).order_by('-timestamp')[:5]

    context = {
        'recent_activities': recent_activities,
        'full_name': request.user.full_name
    }
    return render(request, 'dashboard.html', context)


def landing_view(request):
    """Public landing page."""
    return render(request, 'landing.html')




@login_required
def text2speech_view(request):
    # Initial page load (GET request)
    if request.method == 'GET':
        return render(request, 'text2speech.html')

    # Handle File Upload (POST request)
    if request.method == 'POST' and request.FILES.get('document'):
        document = request.FILES['document']
        fs = FileSystemStorage()
        filename = fs.save(document.name, document)
        file_path = fs.path(filename)

        # 1. SAVE ACTIVITY FOR DASHBOARD
        UserActivity.objects.create(
            user=request.user,
            file_name=document.name,
            tool_used="Text to Speech",
            target_url="/api/classroom/text-to-speech/reader/"
        )

        # 2. SYNC SESSION FOR THE REDIRECT VIEW
        request.session['tts_doc'] = filename

        extracted_text = ""
        ext = document.name.split('.')[-1].lower()

        # 3. EXTRACTION LOGIC (The part that was missing)
        try:
            if ext == 'pdf':
                with open(file_path, 'rb') as pdf_file:
                    reader = PyPDF2.PdfReader(pdf_file)
                    for page in reader.pages:
                        extracted_text += (page.extract_text() or "") + "\n"

            elif ext == 'docx':
                doc = docx.Document(file_path)
                extracted_text = "\n".join([para.text for para in doc.paragraphs])

            elif ext == 'txt':
                # Use open/read to handle large text files correctly
                with open(file_path, 'r', encoding='utf-8') as f:
                    extracted_text = f.read()
            else:
                extracted_text = "Unsupported file format."

        except Exception as e:
            extracted_text = f"Error reading file: {str(e)}"

        # 4. PREPARE SENTENCES
        sentences = re.split(r'(?<=[.!?])\s+', extracted_text.strip())
        sentences = [s for s in sentences if s.strip()]

        return render(request, 'text2speech_reader.html', {
            'filename': document.name,
            'sentences': sentences
        })

    return render(request, 'text2speech.html')

# apps/classroom/views.py

def tts_reader(request):
    """View to handle reading an existing file for Text-to-Speech."""
    # Try to get filename from session (set during upload)
    filename = request.session.get('tts_doc')

    if not filename:
        return redirect('text2speech')

    file_path = os.path.join(settings.MEDIA_ROOT, filename)
    extracted_text = ""
    ext = filename.split('.')[-1].lower()

    try:
        if ext == 'pdf':
            with open(file_path, 'rb') as pdf_file:
                reader = PyPDF2.PdfReader(pdf_file)
                for page in reader.pages:
                    extracted_text += page.extract_text() + "\n"
        elif ext == 'docx':
            doc = docx.Document(file_path)
            extracted_text = "\n".join([para.text for para in doc.paragraphs])
        elif ext == 'txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                extracted_text = f.read()
    except Exception as e:
        extracted_text = f"Error reading file: {str(e)}"

    # Split text into sentences for the TTS reader logic
    sentences = re.split(r'(?<=[.!?])\s+', extracted_text.strip())
    sentences = [s for s in sentences if s.strip()]

    return render(request, 'text2speech_reader.html', {
        'filename': filename,
        'sentences': sentences
    })



def reading_assistant_view(request):
    return render(request, 'readingass.html')

def writing_assistant_view(request):
    return render(request, 'writingass.html')

# --- STUDYBUDDY PLACEHOLDERS ---

from django.http import HttpResponse

def chat_assistant_view(request):
    return HttpResponse("<h1 style='font-family: sans-serif; text-align: center; margin-top: 50px;'>Coming soon....</h1>")

def flashcards_view(request):
    return HttpResponse("<h1 style='font-family: sans-serif; text-align: center; margin-top: 50px;'>Coming soon....</h1>")

def quizzes_view(request):
    return HttpResponse("<h1 style='font-family: sans-serif; text-align: center; margin-top: 50px;'>Coming soon....</h1>")


# --- READING ASSISTANT LOGIC ---
def reading_assistant_upload(request):
    """Handles the file upload and saves it to the media folder."""
    if request.method == 'POST' and request.FILES.get('document'):
        document = request.FILES['document']
        fs = FileSystemStorage()
        filename = fs.save(document.name, document)

        # --- STEP B: CREATE THE LOG RECORD HERE ---
        UserActivity.objects.create(
            user=request.user,
            file_name=document.name,
            tool_used="Reading Assistant",
            # This link takes them directly back to the reader for this document
            target_url=f"/api/classroom/reading-reader/"
        )

        # Store the filename in the session so the reader view can find it
        request.session['current_doc'] = filename
        return redirect('reading_reader')

    return render(request, 'readingass.html')



def reading_reader(request):
    """Extracts text from the saved file and prepares the reader page."""
    filename = request.session.get('current_doc')

    # Redirect if someone tries to access /reader/ without uploading a file
    if not filename:
        return redirect('reading_assistant')

    file_path = os.path.join(settings.MEDIA_ROOT, filename)
    extracted_text = ""

    # 1. CHECK FILE EXTENSION & EXTRACT TEXT
    extension = os.path.splitext(filename)[1].lower()

    try:
        if extension == '.pdf':
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    # Extract text and add a newline to keep paragraphs separate
                    extracted_text += (page.extract_text() or "") + "\n"

        elif extension == '.docx':
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                extracted_text += para.text + "\n"

        elif extension == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                extracted_text = f.read()
    except Exception as e:
        extracted_text = f"Error reading file: {str(e)}"

    # 2. CONTEXT FOR THE AI PERSON
    context = {
        'filename': filename,
        'original_text': extracted_text,
        # The AI person will replace the placeholder below with their API call
        'simplified_text': "AI integration pending. This is where the simplified version will appear."
    }
    return render(request, 'reading_reader.html', context)

# --- API TOOLS ---

def clean_text_api(request):
    """Simple API for cleaning text spacing/capitalization."""
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            text = data.get('text', '')
            cleaned = text.strip().replace("  ", " ").capitalize()
            return JsonResponse({'status': 'success', 'cleaned_text': cleaned})
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)}, status=400)


# apps/classroom/views.py

@login_required
def delete_activity(request, activity_id):
    if request.method == 'POST':
        # 1. Find the activity record
        activity = UserActivity.objects.filter(id=activity_id, user=request.user).first()

        if activity:
            # 2. Path to the physical file in media folder
            # Assuming 'file_name' in your model matches the name in media/
            file_path = os.path.join(settings.MEDIA_ROOT, activity.file_name)

            try:
                # 3. Delete the physical file if it exists
                if os.path.exists(file_path):
                    os.remove(file_path)

                # 4. Delete the database record (This prevents it coming back on reload)
                activity.delete()
                return JsonResponse({'status': 'success'})
            except Exception as e:
                return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

    return JsonResponse({'status': 'error', 'message': 'Invalid request'}, status=400)
